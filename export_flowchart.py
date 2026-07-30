from typing import List, Dict, Optional
import io
import re


# ── helpers ──────────────────────────────────────────────────────

def _mol_svg_inner(mol, size=(300, 200)) -> str:
    from rdkit.Chem import Draw
    try:
        svg = Draw.MolToSVG(mol, size=size)
        inner = re.sub(r'^<svg[^>]*>', '', svg, count=1)
        inner = re.sub(r'</svg>\s*$', '', inner, count=1)
        inner = re.sub(
            r'<rect[^>]*style=\'opacity:1\.0;fill:#FFFFFF;stroke:none\'[^>]*>',
            '', inner, count=1
        )
        return inner.strip()
    except Exception:
        return ""


def _wrap(t: str, n: int = 35) -> str:
    return t if len(t) <= n else t[:n - 3] + "..."


def _mol_info_dict(mol, name: str = "") -> Dict:
    from rdkit import Chem
    from rdkit.Chem import Descriptors
    d = {"name": name}
    try:
        d["smiles"] = Chem.MolToSmiles(mol) if mol else ""
    except Exception:
        d["smiles"] = ""
    try:
        d["formula"] = Descriptors.MolecularFormula(mol) if mol else ""
    except Exception:
        d["formula"] = ""
    try:
        d["mw"] = round(Descriptors.MolWt(mol), 2) if mol else 0.0
    except Exception:
        d["mw"] = 0.0
    try:
        inchi = Chem.MolToInchi(mol) if mol else ""
        d["inchi"] = inchi
    except Exception:
        d["inchi"] = ""
    try:
        inchikey = Chem.MolToInchiKey(mol) if mol else ""
        d["inchikey"] = inchikey
    except Exception:
        d["inchikey"] = ""
    return d


def _cond_lines(conditions: Dict) -> (List[str], List[str]):
    above = []
    below = []
    if conditions.get("solvent"):
        above.append(str(conditions["solvent"]))
    if conditions.get("catalyst"):
        above.append(str(conditions["catalyst"]))
    if conditions.get("reagent"):
        above.append(str(conditions["reagent"]))
    temp = conditions.get("temperature")
    time_val = conditions.get("time")
    parts = []
    if temp is not None:
        parts.append(f"{temp} °C")
    if time_val is not None:
        parts.append(f"{time_val} h")
    if parts:
        above.append(", ".join(parts))
    extra = []
    if conditions.get("pressure"):
        extra.append(f"P = {conditions['pressure']} bar")
    if conditions.get("ph"):
        extra.append(f"pH = {conditions['ph']}")
    if conditions.get("concentration"):
        extra.append(f"[{conditions['concentration']} M]")
    if extra:
        below = extra
    return above, below


# ── SVG constants (publication style) ──────────────────────────────

FONT = "Helvetica, Arial, sans-serif"
MOL_W, MOL_H = 280, 180
ARROW_W = 180
PLUS_W = 36
GAP = 28
MARGIN = 50
TITLE_H = 44
NAME_H = 20
FORMULA_H = 16
MW_H = 16
INFO_H = 8 + NAME_H + FORMULA_H + MW_H
CARD_H = MOL_H + INFO_H


def _compound_block(mol_data, label, x, y, mol_size=(280, 180)):
    lines = []
    mol = mol_data.get("mol")
    name = mol_data.get("name", label)
    formula = mol_data.get("formula", "")
    mw = mol_data.get("mw", "")
    mw_str = f"{mw} g/mol" if mw else ""

    cx = x + mol_size[0] // 2

    if mol:
        inner = _mol_svg_inner(mol, size=mol_size)
        if inner:
            lines.append(
                f'<g transform="translate({x}, {y}) scale(1.0)" overflow="visible">{inner}</g>'
            )

    ny = y + mol_size[1] + 4
    lines.append(
        f'<text x="{cx}" y="{ny}" font-family="{FONT}" font-size="12" font-weight="bold" '
        f'fill="#1a1a1a" text-anchor="middle">{_wrap(name)}</text>'
    )
    if formula:
        lines.append(
            f'<text x="{cx}" y="{ny + NAME_H}" font-family="{FONT}" font-size="10" '
            f'fill="#555555" text-anchor="middle">{formula}</text>'
        )
    if mw_str:
        lines.append(
            f'<text x="{cx}" y="{ny + NAME_H + FORMULA_H}" font-family="{FONT}" font-size="10" '
            f'fill="#777777" text-anchor="middle">{mw_str}</text>'
        )
    return "\n".join(lines)


# ── reaction scheme SVG (publication-quality horizontal layout) ────

def build_reaction_scheme_svg(
    reactants: List[Dict],
    product: Optional[Dict],
    conditions: Dict,
) -> str:
    num = len(reactants)
    cond_above, cond_below = _cond_lines(conditions)

    # determine if we need stacked layout (3+ reactants)
    stack = num >= 3

    if stack:
        mol_w_stack = min(MOL_W, 220)
        mol_h_stack = min(MOL_H, 150)
        row_w = mol_w_stack + ARROW_W + MOL_W + GAP * 3
        col_w = row_w
        n_cols = 1
        if num <= 2:
            pass
        else:
            n_cols = 1
        rows_per_col = num
        scheme_w = col_w
        scheme_h = max(rows_per_col * (mol_h_stack + INFO_H + 12), MOL_H + INFO_H)
        mol_sz = (mol_w_stack, mol_h_stack)
    else:
        scheme_w = num * (MOL_W + GAP + PLUS_W) - PLUS_W + GAP + ARROW_W + GAP + MOL_W
        scheme_h = max(CARD_H, MOL_H + INFO_H)
        mol_sz = (MOL_W, MOL_H)

    canvas_w = int(scheme_w + MARGIN * 2)
    canvas_h = int(MARGIN + TITLE_H + GAP + scheme_h + GAP + MARGIN)

    svg = []
    svg.append(
        f'<svg xmlns="http://www.w3.org/2000/svg" '
        f'width="{canvas_w}" height="{canvas_h}" '
        f'viewBox="0 0 {canvas_w} {canvas_h}">'
    )
    svg.append(f'<rect x="0" y="0" width="{canvas_w}" height="{canvas_h}" fill="#ffffff" rx="0"/>')

    # title
    title_y = MARGIN + 16
    svg.append(
        f'<text x="{canvas_w // 2}" y="{title_y}" font-family="{FONT}" '
        f'font-size="18" font-weight="bold" fill="#1a1a1a" text-anchor="middle">'
        f'Reaction Process Flowchart</text>'
    )

    y0 = MARGIN + TITLE_H + GAP

    if stack:
        # stacked vertical layout for 3+ reactants
        mol_w_use, mol_h_use = mol_sz
        r_per_row = 1
        rows = (num + r_per_row - 1) // r_per_row
        for layout in ["reactants"]:
            pass

        left_block_w = mol_w_use
        total_content_w = left_block_w + GAP + ARROW_W + GAP + MOL_W
        offset_x = (canvas_w - total_content_w) // 2

        # reactants stacked
        for i, r in enumerate(reactants):
            ry = y0 + i * (mol_h_use + INFO_H + 12)
            svg.append(_compound_block(r, f"Reactant {i+1}", offset_x, ry, mol_sz))
            if i < num - 1:
                plus_y = ry + (mol_h_use + INFO_H + 12) // 2
                svg.append(
                    f'<text x="{offset_x + mol_w_use + 8}" y="{plus_y}" '
                    f'font-family="{FONT}" font-size="18" font-weight="300" '
                    f'fill="#555555" text-anchor="middle">+</text>'
                )

        arrow_y = y0 + scheme_h // 2
        arrow_x1 = offset_x + mol_w_use + GAP
        arrow_x2 = arrow_x1 + ARROW_W

        # arrow line
        svg.append(
            f'<line x1="{arrow_x1}" y1="{arrow_y}" x2="{arrow_x2}" y2="{arrow_y}" '
            f'stroke="#333333" stroke-width="2.0"/>'
        )
        # arrowhead
        svg.append(
            f'<polygon points="{arrow_x2 + 10},{arrow_y} {arrow_x2},{arrow_y - 6} '
            f'{arrow_x2},{arrow_y + 6}" fill="#333333"/>'
        )

        # conditions above/below arrow
        if cond_above or cond_below:
            cond_x = (arrow_x1 + arrow_x2) // 2
            for j, cl in enumerate(cond_above):
                svg.append(
                    f'<text x="{cond_x}" y="{arrow_y - 14 - (len(cond_above) - 1 - j) * 13}" '
                    f'font-family="{FONT}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )
            for j, cl in enumerate(cond_below):
                svg.append(
                    f'<text x="{cond_x}" y="{arrow_y + 16 + j * 13}" '
                    f'font-family="{FONT}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )

        # product
        prod_x = arrow_x2 + GAP
        prod_y = y0 + (scheme_h - (MOL_H + INFO_H)) // 2
        if product:
            svg.append(_compound_block(product, "Product", prod_x, prod_y, (MOL_W, MOL_H)))

    else:
        # horizontal layout for 1-2 reactants
        total_w_no_margin = canvas_w - MARGIN * 2
        r_total = num * mol_sz[0] + (num - 1) * (GAP + PLUS_W)
        a_total = ARROW_W
        p_total = mol_sz[0]
        content_w = r_total + GAP + a_total + GAP + p_total
        offset_x = (canvas_w - content_w) // 2

        # reactants
        for i, r in enumerate(reactants):
            rx = offset_x + i * (mol_sz[0] + GAP + PLUS_W)
            ry = y0 + (scheme_h - (mol_sz[1] + INFO_H)) // 2
            svg.append(_compound_block(r, f"Reactant {i+1}", rx, ry, mol_sz))
            if i < num - 1:
                plus_x = rx + mol_sz[0] + GAP // 2
                plus_y = ry + (mol_sz[1] + INFO_H) // 2
                svg.append(
                    f'<text x="{plus_x}" y="{plus_y}" '
                    f'font-family="{FONT}" font-size="20" font-weight="300" '
                    f'fill="#555555" text-anchor="middle">+</text>'
                )

        r_end_x = offset_x + r_total
        arrow_x1 = r_end_x + GAP
        arrow_x2 = arrow_x1 + ARROW_W
        arrow_y = y0 + scheme_h // 2

        # arrow line
        svg.append(
            f'<line x1="{arrow_x1}" y1="{arrow_y}" x2="{arrow_x2}" y2="{arrow_y}" '
            f'stroke="#333333" stroke-width="2.0"/>'
        )
        svg.append(
            f'<polygon points="{arrow_x2 + 10},{arrow_y} {arrow_x2},{arrow_y - 6} '
            f'{arrow_x2},{arrow_y + 6}" fill="#333333"/>'
        )

        # conditions above/below
        if cond_above or cond_below:
            cond_x = (arrow_x1 + arrow_x2) // 2
            for j, cl in enumerate(cond_above):
                svg.append(
                    f'<text x="{cond_x}" y="{arrow_y - 14 - (len(cond_above) - 1 - j) * 13}" '
                    f'font-family="{FONT}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )
            for j, cl in enumerate(cond_below):
                svg.append(
                    f'<text x="{cond_x}" y="{arrow_y + 16 + j * 13}" '
                    f'font-family="{FONT}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )

        # product
        prod_x = arrow_x2 + GAP
        prod_y = y0 + (scheme_h - (MOL_H + INFO_H)) // 2
        if product:
            svg.append(_compound_block(product, "Product", prod_x, prod_y))

    svg.append("</svg>")
    return "\n".join(svg)


# ── individual molecule SVG (publication style) ───────────────────

def build_structure_svg(mol, label: str = "", info: Optional[Dict] = None) -> str:
    if info is None:
        info = _mol_info_dict(mol, label)

    mol_w, mol_h = 400, 280
    panel_x = mol_w + 24
    canvas_w = mol_w + panel_x + 24
    row_h = 18
    canvas_h = max(mol_h + 60, 180)

    lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{canvas_w}" height="{canvas_h}" '
        f'viewBox="0 0 {canvas_w} {canvas_h}">',
        f'<rect x="0" y="0" width="{canvas_w}" height="{canvas_h}" fill="#ffffff" rx="0"/>',
    ]
    inner = _mol_svg_inner(mol, size=(mol_w, mol_h))
    if inner:
        lines.append(
            f'<g transform="translate(8, 8) scale(1.0)" overflow="visible">{inner}</g>'
        )
    else:
        lines.append(
            f'<rect x="8" y="8" width="{mol_w}" height="{mol_h}" rx="0" '
            f'fill="none" stroke="#cccccc" stroke-dasharray="4,3"/>'
            f'<text x="{8 + mol_w // 2}" y="{8 + mol_h // 2}" '
            f'font-family="{FONT}" font-size="14" fill="#999999" '
            f'text-anchor="middle">No structure</text>'
        )

    y0 = 16
    items = [
        ("Name", label or info.get("name", "")),
        ("SMILES", info.get("smiles", "")),
        ("Formula", info.get("formula", "")),
        ("MW", f"{info.get('mw', '')} g/mol"),
        ("InChI", info.get("inchi", "")),
        ("InChIKey", info.get("inchikey", "")),
    ]
    lines.append(
        f'<text x="{panel_x}" y="{y0}" font-family="{FONT}" font-size="14" '
        f'font-weight="bold" fill="#333333">{label}</text>'
    )
    for j, (k, v) in enumerate(items):
        y = y0 + 24 + j * row_h
        lines.append(
            f'<text x="{panel_x}" y="{y}" font-family="{FONT}" font-size="11" '
            f'font-weight="bold" fill="#666666">{k}:</text>'
            f'<text x="{panel_x + 72}" y="{y}" font-family="{FONT}" font-size="11" '
            f'fill="#333333">{_wrap(str(v), 42)}</text>'
        )

    lines.append("</svg>")
    return "\n".join(lines)


# ── SVG bytes ────────────────────────────────────────────────────

def export_svg_bytes(
    reactants, product, conditions,
) -> bytes:
    return build_reaction_scheme_svg(
        reactants, product, conditions
    ).encode("utf-8")


def export_structure_svg_bytes(mol, label="", info=None) -> bytes:
    return build_structure_svg(mol, label, info).encode("utf-8")


# ── PDF export (svglib + reportlab) ───────────────────────────────

A4_W, A4_H = 595.28, 841.89  # points
A4_LANDSCAPE_W, A4_LANDSCAPE_H = A4_H, A4_W


def _svg_to_pdf_bytes(svg_str: str, dpi: int = 300) -> bytes:
    from svglib.svglib import svg2rlg
    from reportlab.graphics import renderPDF
    from reportlab.lib.units import mm
    from reportlab.pdfgen import canvas as pdfcanvas

    drawing = svg2rlg(io.StringIO(svg_str))
    dw, dh = drawing.width, drawing.height

    # page size with 15mm margins
    margin = 15 * mm
    pdf_w = max(dw + 2 * margin, A4_W)
    pdf_h = max(dh + 2 * margin, A4_H)

    # auto-orient: use landscape if content is wider than tall
    if dw > dh:
        page_w, page_h = A4_LANDSCAPE_W, A4_LANDSCAPE_H
    else:
        page_w, page_h = A4_W, A4_H

    usable_w = page_w - 2 * margin
    usable_h = page_h - 2 * margin

    # scale down if needed, never scale up
    scale = min(usable_w / dw, usable_h / dh, 1.0)

    buf = io.BytesIO()
    c = pdfcanvas.Canvas(buf, pagesize=(page_w, page_h))
    c.saveState()

    cx = (page_w - dw * scale) / 2.0
    cy = (page_h - dh * scale) / 2.0

    if scale < 1.0:
        c.translate(cx, cy)
        c.scale(scale, scale)
        renderPDF.draw(drawing, c, 0, 0)
    else:
        renderPDF.draw(drawing, c, cx, cy)

    c.restoreState()
    c.showPage()
    c.save()
    return buf.getvalue()


def export_pdf_bytes(
    reactants, product, conditions,
) -> bytes:
    svg = build_reaction_scheme_svg(reactants, product, conditions)
    return _svg_to_pdf_bytes(svg)


def export_structure_pdf_bytes(mol, label="", info=None) -> bytes:
    svg = build_structure_svg(mol, label, info)
    return _svg_to_pdf_bytes(svg)


# ── PNG export (via pypdfium2) ────────────────────────────────────

def _pdf_to_png_bytes(pdf_bytes: bytes, dpi: int = 300) -> bytes:
    import pypdfium2 as pdfium
    pdf_doc = pdfium.PdfDocument(pdf_bytes)
    page = pdf_doc[0]
    bitmap = page.render(scale=dpi / 72)
    pil_image = bitmap.to_pil()
    buf = io.BytesIO()
    pil_image.save(buf, format="PNG")
    pdf_doc.close()
    return buf.getvalue()


def export_png_bytes(
    reactants, product, conditions,
    dpi=300,
) -> bytes:
    pdf = export_pdf_bytes(reactants, product, conditions)
    return _pdf_to_png_bytes(pdf, dpi)


def export_structure_png_bytes(mol, label="", info=None, dpi=300) -> bytes:
    pdf = export_structure_pdf_bytes(mol, label, info)
    return _pdf_to_png_bytes(pdf, dpi)


# ── JPEG export ──────────────────────────────────────────────────

def export_jpeg_bytes(
    reactants, product, conditions,
    dpi=300,
) -> bytes:
    png = export_png_bytes(reactants, product, conditions, dpi)
    from PIL import Image
    buf = io.BytesIO(png)
    img = Image.open(buf)
    out = io.BytesIO()
    img = img.convert("RGB")
    img.save(out, format="JPEG", quality=95)
    return out.getvalue()


def export_structure_jpeg_bytes(mol, label="", info=None, dpi=300) -> bytes:
    png = export_structure_png_bytes(mol, label, info, dpi)
    from PIL import Image
    buf = io.BytesIO(png)
    img = Image.open(buf)
    out = io.BytesIO()
    img = img.convert("RGB")
    img.save(out, format="JPEG", quality=95)
    return out.getvalue()


# ── DOCX export (editable text + tables, matching reference format) ──

def _set_reference_page(doc):
    """Apply reference document page setup: A4 portrait, tight margins, Times New Roman default."""
    from docx.shared import Pt
    from docx.oxml.ns import qn

    for section in doc.sections:
        section.page_width = 7562850    # 21.01 cm (A4)
        section.page_height = 10699750  # 29.70 cm (A4)
        section.top_margin = 101600     # 0.11 in
        section.bottom_margin = 0
        section.left_margin = 0
        section.right_margin = 89535    # 0.10 in

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)
    style.paragraph_format.space_before = 0
    style.paragraph_format.space_after = 0
    style.paragraph_format.line_spacing = 1.0

    rpr = style.element.find(qn("w:rPr"))
    if rpr is None:
        rpr = style.element.makeelement(qn("w:rPr"), {})
        style.element.append(rpr)


def _add_mol_png_to_run(run, mol, width_cm=4.5):
    """Render an RDKit mol to PNG and add it to a run."""
    from docx.shared import Cm
    from rdkit.Chem import Draw
    try:
        img = Draw.MolToImage(mol, size=(300, 200))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        run.add_picture(buf, width=Cm(width_cm))
    except Exception:
        pass


def _add_run(paragraph, text, bold=False, size=10, font_name="Times New Roman", italic=False):
    """Add a formatted run to a paragraph and return it."""
    from docx.shared import Pt
    run = paragraph.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return run


def _make_reaction_table(doc, reactants, product, conditions):
    """Create a 3-column reaction table for a single stage."""
    from docx.shared import Cm, Pt
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    cond_above, cond_below = _cond_lines(conditions)
    num = len(reactants)
    n_rows = max(num, 1)

    table = doc.add_table(rows=n_rows, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Remove table grid borders
    tbl_pr = table._tbl.tblPr
    tbl_borders = tbl_pr.makeelement(qn("w:tblBorders"), {})
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = tbl_borders.makeelement(qn(f"w:{border_name}"), {})
        border.set(qn("w:val"), "none")
        border.set(qn("w:sz"), "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tbl_borders.append(border)
    tbl_pr.append(tbl_borders)

    # Column widths
    col_widths = [Cm(7.5), Cm(5.0), Cm(7.5)]
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col_widths[i]
            cell.paragraphs[0].paragraph_format.space_before = Pt(0)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)

    # Fill reactants
    for i, r in enumerate(reactants):
        cell = table.cell(i, 0)
        cell.paragraphs[0].alignment = 1  # CENTER
        mol = r.get("mol")
        if mol:
            _add_mol_png_to_run(cell.paragraphs[0].add_run(), mol, width_cm=4.0)
        p2 = cell.add_paragraph()
        p2.alignment = 1
        _add_run(p2, r.get("name", f"Reactant {i+1}"), bold=True, size=10)
        if r.get("formula"):
            p3 = cell.add_paragraph()
            p3.alignment = 1
            _add_run(p3, r.get("formula", ""), size=10)
        if r.get("mw"):
            p4 = cell.add_paragraph()
            p4.alignment = 1
            _add_run(p4, f"{r.get('mw')} g/mol", size=10)

    # Middle column — arrow + conditions
    mid_row = min(num - 1, 0) if num > 0 else 0
    for i in range(n_rows):
        cell = table.cell(i, 1)
        cell.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        if i == mid_row:
            cell.paragraphs[0].alignment = 1
            _add_run(cell.paragraphs[0], "\u2192", bold=True, size=24)
            if cond_above:
                for cl in reversed(cond_above):
                    cp = cell.add_paragraph()
                    cp.alignment = 1
                    _add_run(cp, cl, size=10, italic=True)
            if cond_below:
                for cl in cond_below:
                    cp = cell.add_paragraph()
                    cp.alignment = 1
                    _add_run(cp, cl, size=10, italic=True)

    # Product column
    if product:
        cell = table.cell(0, 2)
        if n_rows > 1:
            cell.merge(table.cell(n_rows - 1, 2))
        cell.paragraphs[0].clear()
        cell.paragraphs[0].alignment = 1
        mol_p = product.get("mol")
        if mol_p:
            _add_mol_png_to_run(cell.paragraphs[0].add_run(), mol_p, width_cm=4.0)
        p2 = cell.add_paragraph()
        p2.alignment = 1
        _add_run(p2, product.get("name", "Product"), bold=True, size=10)
        if product.get("formula"):
            p3 = cell.add_paragraph()
            p3.alignment = 1
            _add_run(p3, product.get("formula", ""), size=10)
        if product.get("mw"):
            p4 = cell.add_paragraph()
            p4.alignment = 1
            _add_run(p4, f"{product.get('mw')} g/mol", size=10)

    return table


def export_docx_bytes(
    reactants, product, conditions,
) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_reference_page(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    _add_run(title, "Reaction Process Flowchart", bold=True, size=11)

    _make_reaction_table(doc, reactants, product, conditions)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ── Multi-stage export ────────────────────────────────────────────

STAGE_GAP = 40
SEP_COLOR = "#cccccc"


def _render_single_svg(
    stage: Dict, title: str, y_offset: int, width: int, font=FONT,
) -> (List[str], int):
    """Render one horizontal reaction scheme at y_offset, return (lines, height_used)."""
    reactants = stage.get("reactants", [])
    product = stage.get("product")
    conditions = stage.get("conditions", {})

    num = len(reactants)
    cond_above, cond_below = _cond_lines(conditions)
    stack = num >= 3

    if stack:
        mw_s, mh_s = 220, 150
        rows_per_col = num
        scheme_h = max(rows_per_col * (mh_s + INFO_H + 12), MOL_H + INFO_H)
        mol_sz = (mw_s, mh_s)
    else:
        scheme_h = max(CARD_H, MOL_H + INFO_H)
        mol_sz = (MOL_W, MOL_H)

    lines = []
    y0 = y_offset

    # stage label
    lines.append(
        f'<text x="{width // 2}" y="{y0}" font-family="{font}" '
        f'font-size="13" font-weight="bold" fill="#333333" text-anchor="middle">{title}</text>'
    )
    y0 += 22

    if stack:
        mw_u, mh_u = mol_sz
        total_cw = mw_u + GAP + ARROW_W + GAP + MOL_W
        ox = (width - total_cw) // 2

        for i, r in enumerate(reactants):
            ry = y0 + i * (mh_u + INFO_H + 12)
            lines.append(_compound_block(r, f"Reactant {i+1}", ox, ry, mol_sz))
            if i < num - 1:
                py = ry + (mh_u + INFO_H + 12) // 2
                lines.append(
                    f'<text x="{ox + mw_u + 8}" y="{py}" '
                    f'font-family="{font}" font-size="18" font-weight="300" '
                    f'fill="#555555" text-anchor="middle">+</text>'
                )

        arrow_y = y0 + scheme_h // 2
        ax1 = ox + mw_u + GAP
        ax2 = ax1 + ARROW_W

        lines.append(f'<line x1="{ax1}" y1="{arrow_y}" x2="{ax2}" y2="{arrow_y}" stroke="#333333" stroke-width="2.0"/>')
        lines.append(
            f'<polygon points="{ax2 + 10},{arrow_y} {ax2},{arrow_y - 6} '
            f'{ax2},{arrow_y + 6}" fill="#333333"/>'
        )
        if cond_above or cond_below:
            cx = (ax1 + ax2) // 2
            for j, cl in enumerate(cond_above):
                lines.append(
                    f'<text x="{cx}" y="{arrow_y - 14 - (len(cond_above) - 1 - j) * 13}" '
                    f'font-family="{font}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )
            for j, cl in enumerate(cond_below):
                lines.append(
                    f'<text x="{cx}" y="{arrow_y + 16 + j * 13}" '
                    f'font-family="{font}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )

        prod_x = ax2 + GAP
        prod_y = y0 + (scheme_h - (MOL_H + INFO_H)) // 2
        if product:
            lines.append(_compound_block(product, "Product", prod_x, prod_y, (MOL_W, MOL_H)))
    else:
        r_total = num * mol_sz[0] + (num - 1) * (GAP + PLUS_W)
        content_w = r_total + GAP + ARROW_W + GAP + mol_sz[0]
        ox = (width - content_w) // 2

        for i, r in enumerate(reactants):
            rx = ox + i * (mol_sz[0] + GAP + PLUS_W)
            ry = y0 + (scheme_h - (mol_sz[1] + INFO_H)) // 2
            lines.append(_compound_block(r, f"Reactant {i+1}", rx, ry, mol_sz))
            if i < num - 1:
                px = rx + mol_sz[0] + GAP // 2
                py = ry + (mol_sz[1] + INFO_H) // 2
                lines.append(
                    f'<text x="{px}" y="{py}" '
                    f'font-family="{font}" font-size="20" font-weight="300" '
                    f'fill="#555555" text-anchor="middle">+</text>'
                )

        r_end_x = ox + r_total
        ax1 = r_end_x + GAP
        ax2 = ax1 + ARROW_W
        arrow_y = y0 + scheme_h // 2

        lines.append(f'<line x1="{ax1}" y1="{arrow_y}" x2="{ax2}" y2="{arrow_y}" stroke="#333333" stroke-width="2.0"/>')
        lines.append(
            f'<polygon points="{ax2 + 10},{arrow_y} {ax2},{arrow_y - 6} '
            f'{ax2},{arrow_y + 6}" fill="#333333"/>'
        )
        if cond_above or cond_below:
            cx = (ax1 + ax2) // 2
            for j, cl in enumerate(cond_above):
                lines.append(
                    f'<text x="{cx}" y="{arrow_y - 14 - (len(cond_above) - 1 - j) * 13}" '
                    f'font-family="{font}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )
            for j, cl in enumerate(cond_below):
                lines.append(
                    f'<text x="{cx}" y="{arrow_y + 16 + j * 13}" '
                    f'font-family="{font}" font-size="10" fill="#444444" text-anchor="middle">{cl}</text>'
                )

        prod_x = ax2 + GAP
        prod_y = y0 + (scheme_h - (MOL_H + INFO_H)) // 2
        if product:
            lines.append(_compound_block(product, "Product", prod_x, prod_y))

    y_use = y0 + scheme_h
    total_h = (y_use + 24) - y_offset
    return lines, total_h


def build_multi_stage_scheme_svg(stages: List[Dict]) -> str:
    num_stages = len(stages)
    if num_stages == 0:
        return '<svg xmlns="http://www.w3.org/2000/svg" width="400" height="200" viewBox="0 0 400 200"><rect width="100%" height="100%" fill="#ffffff"/><text x="200" y="100" font-family="Arial" font-size="14" fill="#999" text-anchor="middle">No stages</text></svg>'

    # Calculate max width needed across all stages
    max_w = MARGIN * 2
    for s in stages:
        reactants = s.get("reactants", [])
        num = len(reactants)
        if num >= 3:
            w = 220 + GAP + ARROW_W + GAP + MOL_W
        else:
            w = num * (MOL_W + GAP + PLUS_W) - PLUS_W + GAP + ARROW_W + GAP + MOL_W
        max_w = max(max_w, w + MARGIN * 2)

    canvas_w = int(max_w)

    # Render each stage, accumulate height
    all_lines = []
    total_h = MARGIN

    # title
    all_lines.append(
        f'<text x="{canvas_w // 2}" y="{MARGIN + 16}" font-family="{FONT}" '
        f'font-size="18" font-weight="bold" fill="#1a1a1a" text-anchor="middle">'
        f'Reaction Process Flowchart</text>'
    )
    total_h += TITLE_H + 8

    for i, s in enumerate(stages):
        # separator line
        if i > 0:
            sep_y = total_h + STAGE_GAP // 2
            all_lines.append(
                f'<line x1="{MARGIN}" y1="{sep_y}" x2="{canvas_w - MARGIN}" y2="{sep_y}" '
                f'stroke="{SEP_COLOR}" stroke-width="1.0" stroke-dasharray="6,4"/>'
            )
            total_h += STAGE_GAP

        title = f"Stage {s.get('stage_number', i + 1)} of {num_stages}"
        stage_lines, used = _render_single_svg(s, title, total_h, canvas_w)
        all_lines.extend(stage_lines)
        total_h += used + 8

    total_h += MARGIN

    svg_lines = [
        f'<svg xmlns="http://www.w3.org/2000/svg" '
        f'width="{canvas_w}" height="{int(total_h)}" '
        f'viewBox="0 0 {canvas_w} {int(total_h)}">',
        f'<rect x="0" y="0" width="{canvas_w}" height="{int(total_h)}" fill="#ffffff" rx="0"/>',
    ]
    svg_lines.extend(all_lines)
    svg_lines.append("</svg>")
    return "\n".join(svg_lines)


def export_multi_stage_svg_bytes(stages: List[Dict]) -> bytes:
    return build_multi_stage_scheme_svg(stages).encode("utf-8")


def export_multi_stage_pdf_bytes(stages: List[Dict]) -> bytes:
    svg = build_multi_stage_scheme_svg(stages)
    return _svg_to_pdf_bytes(svg)


def export_multi_stage_png_bytes(stages: List[Dict], dpi=300) -> bytes:
    pdf = export_multi_stage_pdf_bytes(stages)
    return _pdf_to_png_bytes(pdf, dpi)


def export_multi_stage_jpeg_bytes(stages: List[Dict], dpi=300) -> bytes:
    png = export_multi_stage_png_bytes(stages, dpi)
    from PIL import Image
    buf = io.BytesIO(png)
    img = Image.open(buf)
    out = io.BytesIO()
    img = img.convert("RGB")
    img.save(out, format="JPEG", quality=95)
    return out.getvalue()


def export_multi_stage_docx_bytes(stages: List[Dict]) -> bytes:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_reference_page(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(12)
    _add_run(title, "Reaction Process Flowchart", bold=True, size=11)

    for i, s in enumerate(stages):
        if i > 0:
            doc.add_paragraph()

        stage_title = doc.add_paragraph()
        stage_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        stage_title.paragraph_format.space_after = Pt(12)
        _add_run(stage_title, f"Stage {s.get('stage_number', i + 1)} of {len(stages)}", bold=True, size=10)

        reactants = s.get("reactants", [])
        product = s.get("product")
        conditions = s.get("conditions", {})
        _make_reaction_table(doc, reactants, product, conditions)

        doc.add_paragraph()  # one line gap after each stage

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def embed_mol_image(doc, mol, width_inches=3.0):
    from docx.shared import Inches
    from rdkit.Chem import Draw
    try:
        img = Draw.MolToImage(mol, size=(400, 280))
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        buf.seek(0)
        p = doc.add_paragraph()
        p.alignment = 1
        p.add_run().add_picture(buf, width=Inches(width_inches))
    except Exception:
        pass


def export_structure_docx_bytes(mol, label="", info=None) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if info is None:
        info = _mol_info_dict(mol, label)

    doc = Document()
    title = doc.add_heading(label or "Molecular Structure", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    embed_mol_image(doc, mol, width_inches=5.0)

    items = [
        ("Name", label or info.get("name", "")),
        ("SMILES", info.get("smiles", "")),
        ("Formula", info.get("formula", "")),
        ("Molecular Weight", f"{info.get('mw', '')} g/mol"),
        ("InChI", info.get("inchi", "")),
        ("InChIKey", info.get("inchikey", "")),
    ]
    table = doc.add_table(rows=len(items), cols=2)
    table.style = "Light Shading Accent 1"
    for j, (k, v) in enumerate(items):
        cell_k = table.cell(j, 0)
        cell_k.text = k
        for paragraph in cell_k.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        table.cell(j, 1).text = str(v)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
