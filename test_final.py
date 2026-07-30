from rdkit import Chem
from export_flowchart import export_docx_bytes, export_multi_stage_docx_bytes
from docx import Document
from docx.oxml.ns import qn
import io

mol1 = Chem.MolFromSmiles("CCO")
mol2 = Chem.MolFromSmiles("CC=O")
reacts = [{"name":"Ethanol","smiles":"CCO","formula":"C2H6O","mw":46.07,"mol":mol1}]
prods = {"name":"Acetaldehyde","smiles":"CC=O","formula":"C2H4O","mw":44.05,"mol":mol2}
conds = {"solvent":"water","catalyst":"","reagent":"NaBH4","temperature":25,"time":2}

# Multi-stage
b = export_multi_stage_docx_bytes([
    {"stage_number":1,"reactants":reacts,"product":prods,"conditions":conds},
    {"stage_number":2,"reactants":reacts,"product":prods,"conditions":conds},
])
doc = Document(io.BytesIO(b))
print("=== Multi-stage ===")
for pi, p in enumerate(doc.paragraphs):
    txt = p.text[:60]
    space = p.paragraph_format.space_after
    if txt.strip():
        sz = p.runs[0].font.size / 12700 if p.runs else "?"
        print(f"P{pi}: sz={sz}pt space_after={space} txt={repr(txt)}")
    else:
        print(f"P{pi}: (gap) space_after={space}")

print("\nImage sizes:")
for ti, t in enumerate(doc.tables):
    for ri, row in enumerate(t.rows):
        for ci, cell in enumerate(row.cells):
            extents = cell._element.findall(".//" + qn("wp:extent"))
            for ext in extents:
                cx = int(ext.get("cx"))
                in_cm = cx / 360000
                print(f"  T{ti}[{ri},{ci}] {in_cm:.1f} cm")

# Single-stage
print("\n=== Single-stage ===")
b1 = export_docx_bytes(reacts, prods, conds)
doc1 = Document(io.BytesIO(b1))
for pi, p in enumerate(doc1.paragraphs):
    txt = p.text[:60]
    if txt.strip():
        sz = p.runs[0].font.size / 12700 if p.runs else "?"
        space = p.paragraph_format.space_after
        print(f"P{pi}: sz={sz}pt space_after={space} txt={repr(txt)}")

tables = doc.tables
assert len(tables) == 2
assert doc.sections[0].page_width == 7562850
assert doc.styles["Normal"].font.name == "Times New Roman"
print("\nAll OK")
